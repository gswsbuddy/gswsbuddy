from http.server import BaseHTTPRequestHandler
import json
from datetime import datetime, timedelta
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        content_length = int(self.headers['Content-Length'])
        post_data = self.rfile.read(content_length)
        
        try:
            data = json.loads(post_data.decode('utf-8'))
        except Exception:
            self.send_response(400)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({"error": "Invalid JSON payload structure."}).encode())
            return

        # 1. Gather properties out of payload dictionary
        scope = data.get('scope', 'RURAL').upper()
        name = data.get('name', '').strip() or "[Name Placeholder]"
        gender = data.get('gender', 'Male')
        designation = data.get('designation', '').strip() or "[Designation]"
        place_of_working = data.get('placeOfWorking', '').strip() or "[Secretariat]"
        mandal = data.get('mandal', '').strip() or "[Mandal]"
        district = data.get('district', '').strip() or "[District]"
        
        # Date Processing Setup
        joining_raw = data.get('dateOfJoining', '')
        reg_raw = data.get('regularisationDate', '')
        inc_raw = data.get('lastIncrementDate', '')
        leaves_input = data.get('leaves', {})
        
        def format_date(d_str):
            try:
                return datetime.strptime(d_str, '%Y-%m-%d').strftime('%d/%m/%Y')
            except:
                return "NIL"

        date_of_joining = format_date(joining_raw)
        regularisation_date = format_date(reg_raw)
        last_increment_date = format_date(inc_raw)
        
        # Calculate dynamic loop date ceiling (+364 days offset)
        one_year_date = "[Date]"
        if inc_raw:
            try:
                dt = datetime.strptime(inc_raw, '%Y-%m-%d') + timedelta(days=364)
                one_year_date = dt.strftime('%d/%m/%Y')
            except:
                pass

        basic_pay = data.get('basicPay', '') or "[Amount]"

        # 2. Assign dynamic contextual naming parameters
        branding_label = "Swarna Gramam" if scope == "RURAL" else "Swarna Ward"
        authority_designation = "Panchayath Secretary (DDO)" if scope == "RURAL" else "Ward Administrative Secretary"

        title_prefix = "Sri" if gender == "Male" else "Smt"
        he_she = "He" if gender == "Male" else "She"
        his_her = "His" if gender == "Male" else "Her"

        # 3. Document Building & Formatting Options
        doc = Document()
        
        # Apply standard page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # Apply Global Font Styling Rules (Times New Roman, 12pt)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Document Header Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("EMPLOYEE SERVICE CERTIFICATE")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.underline = True
        title_p.paragraph_format.space_after = Pt(24)

        # Main Certificate Body Narrative
        body_p = doc.add_paragraph()
        body_p.paragraph_format.line_spacing = 1.5
        body_p.paragraph_format.space_after = Pt(12)
        body_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_p.paragraph_format.first_line_indent = Inches(0.5)
        
        narrative_text = (
            f"This is to certify that {title_prefix} {name} working as {designation} "
            f"at {place_of_working} {branding_label}, {mandal} (M), {district} District "
            f"from {date_of_joining}. {he_she} has regularised from service on {regularisation_date}. "
            f"{his_her} basic pay is Rs. {basic_pay}/- wef {last_increment_date}. "
            f"{he_she} has completed one year of {his_her.lower()} service without any break. "
            f"{he_she} has paid and drawn salary for the period from {last_increment_date} to {one_year_date}."
        )
        body_p.add_run(narrative_text)

        # Table Introduction
        intro_p = doc.add_paragraph()
        intro_p.paragraph_format.line_spacing = 1.5
        intro_p.paragraph_format.space_after = Pt(6)
        intro_p.add_run(f"{his_her} leave availed details are as follows from the period {last_increment_date} to {one_year_date}:")

        # 4. Generate Leave Tracking Matrix Table Grid
        table = doc.add_table(rows=6, cols=6)
        table.style = 'Table Grid'
        
        headers = ["S. No", "Type of Leave", "From", "To", "Days", "Remarks"]
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(headers):
            hdr_cells[i].text = text
            p = hdr_cells[i].paragraphs[0]
            p.runs[0].font.bold = True
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            if i != 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        leave_keys = [
            ("maternity", "Maternity Leave"),
            ("eol", "EOL - ExtraOrdinary Leave"),
            ("medical", "Medical Leave (HPL/Commuted)"),
            ("childcare", "Child Care Leave"),
            ("earned", "Earned Leave")
        ]
        
        def calc_days(start_s, end_s):
            if not start_s or not end_s: return 0
            try:
                s = datetime.strptime(start_s, '%Y-%m-%d')
                e = datetime.strptime(end_s, '%Y-%m-%d')
                delta = (e - s).days + 1
                return delta if delta > 0 else 0
            except:
                return 0

        for idx, (key, label) in enumerate(leave_keys, start=1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text = label
            
            leave_data = leaves_input.get(key, {})
            is_availed = leave_data.get('availed', False)
            
            if is_availed:
                f_date = format_date(leave_data.get('from', ''))
                t_date = format_date(leave_data.get('to', ''))
                days = calc_days(leave_data.get('from', ''), leave_data.get('to', ''))
                remarks = leave_data.get('remarks', '').strip() or "-"
                
                row_cells[2].text = f_date
                row_cells[3].text = t_date
                row_cells[4].text = str(days)
                row_cells[5].text = remarks
            else:
                row_cells[2].text = "NIL"
                row_cells[3].text = "NIL"
                row_cells[4].text = "NIL"
                row_cells[5].text = "-"

            for col_idx in [0, 2, 3, 4, 5]:
                p = row_cells[col_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)

        # Eligibility Declaration Summary Block Paragraph
        eligibility_p = doc.add_paragraph()
        eligibility_p.paragraph_format.space_before = Pt(12)
        eligibility_p.paragraph_format.line_spacing = 1.5
        eligibility_p.add_run(f"So {title_prefix} {name}, {designation} is eligible for Annual Grade Increment wef {one_year_date} as per our records.")

        # 5. Right-Aligned Issuing Authority Signature Block
        sig_p = doc.add_paragraph()
        sig_p.paragraph_format.space_before = Pt(40)
        sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_run = sig_p.add_run(
            f"{authority_designation}\n"
            f"{place_of_working} {branding_label}\n"
            f"{mandal} (M.), {district} Dist."
        )
        sig_run.font.bold = True

        # 6. Stream file back down memory container pipes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        binary_output = file_stream.read()

        self.send_response(200)
        self.send_header('Content-type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        self.send_header('Content-Disposition', f'attachment; filename="{name}_Increment_Certificate.docx"')
        self.send_header('Content-Length', str(len(binary_output)))
        self.end_headers()
        self.wfile.write(binary_output)