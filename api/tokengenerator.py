from http.server import BaseHTTPRequestHandler
import json
from urllib.parse import urlparse, parse_qs
from datetime import datetime
import io

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class handler(BaseHTTPRequestHandler):
    def do_GET(self):
        query = urlparse(self.path).query
        params = parse_qs(query)
        
        rsk = params.get('rsk', ['Unknown RBK'])[0].strip()
        mandal = params.get('mandal', ['Unknown Mandal'])[0].strip()
        raw_date = params.get('date', [''])[0].strip()
        
        try:
            token_count = int(params.get('count', [100])[0])
            if token_count > 500: token_count = 500
        except (ValueError, TypeError):
            token_count = 100

        try:
            date_obj = datetime.strptime(raw_date, "%Y-%m-%d")
            formatted_date = date_obj.strftime("%d-%b-%Y")
        except Exception:
            formatted_date = raw_date

        doc = Document()
        
        # 1. Strict Page Alignment (A4 dimensions with 0.5-inch margins)
        for section in doc.sections:
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        token_index = 1

        while token_index <= token_count:
            # 2. Generate a perfectly dimensioned 4x3 Table Grid
            table = doc.add_table(rows=4, cols=3)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.autofit = False
            
            # Apply standard borders to table XML so cutting guidelines are visible
            tblPr = table._tbl.tblPr
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>\n'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
                f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
                f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
                f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
                f'</w:tblBorders>'
            )
            tblPr.append(borders)

            for row_idx in range(4):
                row = table.rows[row_idx]
                row.height = Inches(2.6) # Strict fixed token height target to distribute 4 rows on a page
                
                for col_idx in range(3):
                    cell = row.cells[col_idx]
                    cell.width = Inches(2.42)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    if token_index <= token_count:
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # CRITICAL: Strip out Word's automatic spacing gaps entirely
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.15

                        # Line 1: Header Context
                        r1 = p.add_run(f"Urea Distribution – {formatted_date}\n")
                        r1.font.name = 'ArialBlack'
                        r1.font.size = Pt(11)
                        r1.font.bold = True
                        r1.font.color.rgb = None

                        # Line 2: Giant Token ID
                        r2 = p.add_run(f"\n{token_index}\n\n")
                        r2.font.name = 'Arial'
                        r2.font.size = Pt(28)
                        r2.font.bold = True

                        # Line 3: Validation baseline
                        r3 = p.add_run("Signature of VAA / MAO\n")
                        r3.font.name = 'Arial'
                        r3.font.size = Pt(9)
                        r3.font.bold = True

                        # Line 4: Organization Tracker Localizer
                        r4 = p.add_run(f"{rsk} RSK | {mandal} Mdl")
                        r4.font.name = 'Arial'
                        r4.font.size = Pt(9)
                        r4.font.bold = True
                        
                        token_index += 1
                    else:
                        p = cell.paragraphs[0]
                        p.text = ""
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)

            # Prevent trailing empty pages on the last loop pass
            if token_index <= token_count:
                doc.add_page_break()

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        binary_data = file_stream.read()

        clean_filename = f"{rsk.replace(' ', '_')}_Tokens.docx"
        self.send_response(200)
        self.send_header('Content-type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        self.send_header('Content-Disposition', f'attachment; filename="{clean_filename}"')
        self.end_headers()
        self.wfile.write(binary_data)